#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from pathlib import Path
import math
import re

from PIL import Image, ImageDraw
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "实验6_控制实验报告.md"
PDF_PATH = ROOT / "实验6_控制实验报告.pdf"
DOCX_PATH = ROOT / "实验6_控制实验报告.docx"
IMG_PATH = ROOT / "results" / "pid_tracking_curve.png"


def load_points(path):
    rows = []
    for line in path.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        x, y, *_ = [float(v) for v in line.split(",")]
        rows.append((x, y))
    return rows


def simulate_pid(reference):
    x, y, yaw = 4.8, -0.8, math.pi / 2
    v, dt, wheel_base = 0.45, 0.05, 0.313
    kp, ki, kd = 0.85, 0.02, 0.18
    integral, last_error = 0.0, 0.0
    trajectory = []

    for _ in range(900):
        nearest = min(range(len(reference)), key=lambda i: (reference[i][0] - x) ** 2 + (reference[i][1] - y) ** 2)
        target = reference[min(nearest + 6, len(reference) - 1)]
        dx, dy = target[0] - x, target[1] - y
        lateral_error = -math.sin(yaw) * dx + math.cos(yaw) * dy
        integral += lateral_error * dt
        derivative = (lateral_error - last_error) / dt
        steering = kp * lateral_error + ki * integral + kd * derivative
        steering = max(-0.523599, min(0.523599, steering))
        last_error = lateral_error

        omega = v * math.tan(steering) / wheel_base
        dtheta = omega * dt
        x += v * math.cos(yaw + dtheta / 2) * dt
        y += v * math.sin(yaw + dtheta / 2) * dt
        yaw = (yaw + dtheta + math.pi) % (2 * math.pi) - math.pi
        trajectory.append((x, y))

    return trajectory


def draw_tracking_plot():
    reference = load_points(ROOT / "source" / "pid_controller" / "waypoints" / "circle.csv")
    trajectory = simulate_pid(reference)
    points = reference + trajectory
    min_x, max_x = min(p[0] for p in points), max(p[0] for p in points)
    min_y, max_y = min(p[1] for p in points), max(p[1] for p in points)

    width, height, margin = 1100, 760, 80
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    def tx(point):
        x, y = point
        sx = margin + (x - min_x) / (max_x - min_x) * (width - 2 * margin)
        sy = height - margin - (y - min_y) / (max_y - min_y) * (height - 2 * margin)
        return sx, sy

    for gx in range(math.floor(min_x), math.ceil(max_x) + 1):
        x0, _ = tx((gx, min_y))
        draw.line((x0, margin, x0, height - margin), fill=(235, 238, 240))
    for gy in range(math.floor(min_y), math.ceil(max_y) + 1):
        _, y0 = tx((min_x, gy))
        draw.line((margin, y0, width - margin, y0), fill=(235, 238, 240))

    draw.line([tx(p) for p in reference], fill=(32, 96, 180), width=5)
    draw.line([tx(p) for p in trajectory], fill=(210, 62, 62), width=4)
    draw.rectangle((margin, margin, width - margin, height - margin), outline=(90, 90, 90), width=2)
    draw.text((margin, 28), "PID path tracking: reference path (blue) vs tracked trajectory (red)", fill=(30, 30, 30))
    draw.text((margin, height - 48), "Kp=0.85  Ki=0.02  Kd=0.18  v=0.45m/s  steering limit=+-30deg", fill=(60, 60, 60))
    IMG_PATH.parent.mkdir(parents=True, exist_ok=True)
    img.save(IMG_PATH)


def markdown_to_flowables(markdown):
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    normal = ParagraphStyle("cn-normal", parent=styles["Normal"], fontName="STSong-Light", fontSize=10.5, leading=17)
    h1 = ParagraphStyle("cn-h1", parent=normal, fontSize=18, leading=24, textColor=colors.HexColor("#285a8c"), spaceAfter=10)
    h2 = ParagraphStyle("cn-h2", parent=normal, fontSize=14, leading=20, textColor=colors.HexColor("#285a8c"), spaceBefore=10, spaceAfter=6)
    h3 = ParagraphStyle("cn-h3", parent=normal, fontSize=12, leading=18, textColor=colors.HexColor("#285a8c"), spaceBefore=8, spaceAfter=4)
    code = ParagraphStyle("cn-code", parent=normal, fontName="Courier", fontSize=8.5, leading=12, backColor=colors.HexColor("#f5f5f2"), leftIndent=8, rightIndent=8)

    flowables = []
    in_code = False
    code_lines = []

    def flush_code():
        if code_lines:
            text = "<br/>".join(line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;") for line in code_lines)
            flowables.append(Paragraph(text, code))
            flowables.append(Spacer(1, 0.12 * cm))
            code_lines.clear()

    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line:
            flowables.append(Spacer(1, 0.12 * cm))
            continue
        if line.startswith("# "):
            flowables.append(Paragraph(line[2:], h1))
        elif line.startswith("## "):
            flowables.append(Paragraph(line[3:], h2))
        elif line.startswith("### "):
            flowables.append(Paragraph(line[4:], h3))
        elif line.startswith("!["):
            match = re.search(r"\(([^)]+)\)", line)
            if match:
                image_path = ROOT / match.group(1)
                flowables.append(RLImage(str(image_path), width=15.8 * cm, height=10.9 * cm))
                flowables.append(Spacer(1, 0.2 * cm))
        elif line.startswith("- "):
            flowables.append(Paragraph("• " + line[2:], normal))
        else:
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            safe = re.sub(r"`([^`]+)`", r"<font name='Courier'>\1</font>", safe)
            flowables.append(Paragraph(safe, normal))

    flush_code()
    return flowables


def build_pdf():
    markdown = MD_PATH.read_text(encoding="utf-8")
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=2.1 * cm,
        rightMargin=2.1 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
        title="控制实验报告",
    )
    doc.build(markdown_to_flowables(markdown))


def build_docx():
    markdown = MD_PATH.read_text(encoding="utf-8")
    document = Document()
    document.styles["Normal"].font.name = "宋体"
    document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    in_code = False
    code_lines = []
    for line in markdown.splitlines():
        if line.startswith("```"):
            if in_code:
                document.add_paragraph("\n".join(code_lines), style="Intense Quote")
                code_lines.clear()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=2)
        elif line.startswith("!["):
            document.add_picture(str(IMG_PATH), width=Inches(6.2))
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            document.add_paragraph(line)
        else:
            document.add_paragraph("")
    document.save(DOCX_PATH)


def main():
    draw_tracking_plot()
    build_pdf()
    build_docx()
    print(PDF_PATH)
    print(DOCX_PATH)
    print(IMG_PATH)


if __name__ == "__main__":
    main()
